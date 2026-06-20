"""Verify the final submission docx."""
from pathlib import Path
from docx import Document
import re

REPO = Path(__file__).resolve().parent.parent
OUT = REPO / "#GRADUATION_FINAL" / "CEREBRO_Graduation_Project_FINAL_SUBMISSION.docx"

d = Document(str(OUT))

print("=" * 60)
print("CEREBRO Final Submission DOCX - Verification")
print("=" * 60)
print(f"File: {OUT.name}")
print(f"Size: {OUT.stat().st_size / 1024 / 1024:.1f} MB")
print(f"Paragraphs: {len(d.paragraphs)}")
print(f"Tables: {len(d.tables)}")
print(f"Inline shapes: {len(d.inline_shapes)}")

# Include table cell contents so team names + competitor tables are searchable.
table_text_parts = []
for t in d.tables:
    for r in t.rows:
        for c in r.cells:
            table_text_parts.append(c.text)
text = "\n".join(p.text for p in d.paragraphs) + "\n" + "\n".join(table_text_parts)
refs = re.findall(r"^\[(\d+)\]", text, re.M)
n_refs = max((int(r) for r in refs), default=0)
print(f"Max reference index: {n_refs}")
print(f"Approx word count: {sum(len(p.text.split()) for p in d.paragraphs)}")

print()
print("--- Critical sections present? ---")
expected = [
    "Business Model Canvas",
    "Software-Stack and Niche Competitor",
    "Table 2.3",
    "Table 2.4",
    "Appendix G",
    "Project Photo Gallery",
    "LIST OF FIGURES",
    "Ahmed Mohamed Moussa",
    "Sandy Samy Samir",
    "Basma Ahmed Elmorsy",
    "222101392",
    "222101524",
    "221101164",
    "Aya Zoghby",
    "Khaled Fouad",
    "Moawad El-Kholy",
    "Brilliant Labs",
    "OrCam MyEye",
    "MentraOS",
    "Even Realities",
    "Biped",
    ".Lumen",
    "Envision Glasses",
    "CEREBRO Business Model Canvas",
]
for e in expected:
    found = e in text
    flag = "OK" if found else "MISS"
    print(f"  [{flag}] {e}")

print()
print("--- Spelling sweep ---")
bad = ["researchh", "ferroferroconcrete", "researc "]
for b in bad:
    occ = text.count(b)
    print(f"  '{b}' occurrences: {occ}")

print()
print("--- BMC table summary ---")
for i, t in enumerate(d.tables):
    first = t.rows[0].cells[0].text[:40] if t.rows else ""
    if "Key Partners" in first or "Key Partners" in t.rows[0].cells[0].text:
        print(f"  Table {i}: {len(t.rows)}x{len(t.columns)}")
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                first_line = cell.text.split("\n")[0][:35]
                print(f"    ({ri},{ci}) {first_line}")
        break

print()
print("--- Competitor tables summary ---")
for i, t in enumerate(d.tables):
    first = t.rows[0].cells[0].text if t.rows else ""
    if first in ("Platform / Stack", "Wearable / platform"):
        print(f"  Table {i} ({first}): {len(t.rows)}x{len(t.columns)} rows")

print()
print("--- Last few inline figure parents ---")
print(f"  Total inline shapes: {len(d.inline_shapes)}")
print("  (10 chapter figures + 135 gallery thumbnails expected)")

print()
print("--- References preview ---")
ref_lines = [p.text for p in d.paragraphs if re.match(r"^\[\d+\]", p.text)]
print(f"  Total reference lines: {len(ref_lines)}")
print(f"  First: {ref_lines[0][:80] if ref_lines else 'NONE'}")
print(f"  Last:  {ref_lines[-1][:80] if ref_lines else 'NONE'}")
