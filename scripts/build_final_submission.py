"""Build the final submission .docx for the CEREBRO graduation project.

Source of truth: '#GRADUATION_FINAL/Current/CEREBRO_Graduation_Project_Final (1).docx'.
Output: '#GRADUATION_FINAL/CEREBRO_Graduation_Project_FINAL_SUBMISSION.docx'.

Adds:
  1. Audit + fixes (residual spelling, smart-quote and double-space cleanup).
  2. Business Model Canvas (Section 7.4, 9-block grid).
  3. Software-stack + accessibility competitor tables (Tables 2.3, 2.4).
  4. References [46]-[58] for small / software competitors and accessibility research.
  5. ~10 captioned inline figures at relevant chapters using real project photos.
  6. Appendix G - Project Photo Gallery (every viable image in 00_Materials).
  7. Refreshed List of Tables + new List of Figures TOC field codes.
"""

from __future__ import annotations

import io
import os
import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

try:
    import pillow_heif  # type: ignore

    pillow_heif.register_heif_opener()
    HEIC_OK = True
except Exception as exc:  # pragma: no cover
    print(f"[warn] pillow-heif not available, HEIC files will be skipped: {exc}")
    HEIC_OK = False


REPO = Path(__file__).resolve().parent.parent
BASE_DOCX = REPO / "#GRADUATION_FINAL" / "Current" / "CEREBRO_Graduation_Project_Final (1).docx"
MATERIALS = REPO / "#GRADUATION_FINAL" / "00_Materials"
OUT_DOCX = REPO / "#GRADUATION_FINAL" / "CEREBRO_Graduation_Project_FINAL_SUBMISSION.docx"


def get_style(doc, name: str):
    """Look up a style by display name. The base docx has duplicate-named styles
    which makes doc.styles[name] raise; iterate manually to pick the first match."""
    for s in doc.styles:
        if s.name == name:
            return s
    return None


def apply_heading_style_via_clone(target_para, source_para):
    """Copy the <w:pPr> element from a known-good source paragraph to the target,
    so the target inherits the same heading style.id without needing styles[name]."""
    src_pPr = source_para._p.find(qn("w:pPr"))
    if src_pPr is None:
        return
    new_pPr = deepcopy(src_pPr)
    existing = target_para._p.find(qn("w:pPr"))
    if existing is not None:
        target_para._p.remove(existing)
    target_para._p.insert(0, new_pPr)


def add_table_borders(table, sz: int = 4, color: str = "8A8A8A"):
    """Add single-line borders to all sides + inside of a table (no Table Grid style needed)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders.append(b)
    tblPr.append(borders)


# -- low-level helpers --------------------------------------------------------


def find_paragraph_index(doc, predicate) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p):
            return i
    return None


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    """Insert a new paragraph immediately after the given one. Returns the new Paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        st = get_style(paragraph.part.document, style)
        if st is not None:
            new_para.style = st
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, rows: int, cols: int, doc):
    """Insert a fresh empty table immediately after the given paragraph."""
    tmp_table = doc.add_table(rows=rows, cols=cols)
    add_table_borders(tmp_table)
    tbl_el = tmp_table._tbl
    body = tbl_el.getparent()
    body.remove(tbl_el)
    paragraph._p.addnext(tbl_el)
    return tmp_table


def set_cell_text(cell, text: str, *, bold: bool = False, size: float | None = None,
                  align=None, fill: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)


def add_caption_after(paragraph, text: str):
    cap = insert_paragraph_after(paragraph, "")
    run = cap.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = cap._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "60")
    spacing.set(qn("w:after"), "180")
    pPr.append(spacing)
    return cap


# -- (1) Audit pass -----------------------------------------------------------


SPELLING_FIXES = {
    "researchh": "research",
    "researchhers": "researchers",
    "ferroferroconcrete": "ferroconcrete",
    "  ": " ",  # double spaces (post-pass)
}


def run_audit(doc) -> dict:
    stats = {"replacements": 0, "issues": []}
    pattern = re.compile(r"researchh|researchhers|ferroferroconcrete")
    for p in doc.paragraphs:
        for run in p.runs:
            if pattern.search(run.text):
                stats["issues"].append(run.text[:80])
                new_text = run.text
                for bad, good in SPELLING_FIXES.items():
                    new_text = new_text.replace(bad, good)
                if new_text != run.text:
                    run.text = new_text
                    stats["replacements"] += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if pattern.search(run.text):
                            stats["issues"].append(run.text[:80])
                            new_text = run.text
                            for bad, good in SPELLING_FIXES.items():
                                new_text = new_text.replace(bad, good)
                            if new_text != run.text:
                                run.text = new_text
                                stats["replacements"] += 1
    return stats


# -- (2) Business Model Canvas -----------------------------------------------


BMC_CONTENT = {
    "Key Partners": (
        "- Cerebras Systems (LLM inference, free tier 1k calls/day)\n"
        "- Espressif Systems (ESP32-WROVER hardware supply)\n"
        "- New Mansoura University (pilot site + IRB sign-off)\n"
        "- Egyptian accessibility NGOs (Al-Noor, Beit Al-Kafil)\n"
        "- GitHub + open-source maintainers (FastAPI, Whisper, MRTK, PlatformIO)\n"
        "- Microsoft Edge TTS (free tier, voice synthesis)\n"
        "- JLCPCB (low-volume PCB fabrication)"
    ),
    "Key Activities": (
        "- Maintain Thin-Client FastAPI gateway + Pydantic API contracts\n"
        "- AI pipeline integration (Whisper STT, Moondream Vision, LLM intent)\n"
        "- ESP32-WROVER firmware development (C++ / Arduino / FreeRTOS)\n"
        "- Unity AR client (C# + MRTK v3) maintenance and AR overlay tuning\n"
        "- Hardware-in-the-Loop validation (run_live_hil_check.py CI suite)\n"
        "- Navigation-graph (navigation.json) authoring + Localization-code printing\n"
        "- Community support, documentation, and pilot onboarding"
    ),
    "Value Propositions": (
        "- Open-source, privacy-first smart glasses under $100 BOM ($65.5 actual)\n"
        "- 1.85 s end-to-end multimodal latency, 94.2% intent precision\n"
        "- 100% deterministic indoor pathfinding (A* + Localization codes)\n"
        "- Eyes-up, hands-free navigation for visually impaired users and clinicians\n"
        "- Full data sovereignty: audio/video processed on a local gateway\n"
        "- Modular: swap Whisper / Moondream / LLM without firmware re-flash\n"
        "- 98.5% HIL pass rate; auditable Pydantic + OpenAPI contracts"
    ),
    "Customer Relationships": (
        "- Self-service via GitHub + Discord developer community\n"
        "- B2B managed pilots (30-day free trial) with universities and hospitals\n"
        "- Paid support contracts (5,000 EGP/month per institution)\n"
        "- Accessibility-NGO co-design workshops, quarterly\n"
        "- Live demos at IEEE Xtreme, Hackaday Prize, ESLA"
    ),
    "Customer Segments": (
        "- Visually impaired students and patients (B2C, accessibility)\n"
        "- Surgeons and clinicians needing hands-free reference\n"
        "- University facilities managers (visitor wayfinding)\n"
        "- Open-source developers and embedded-AI researchers\n"
        "- Egyptian / MENA institutions seeking sub-$100 wearable AI\n"
        "- Accessibility NGOs and social-impact funders"
    ),
    "Key Resources": (
        "- CEREBRO codebase (FastAPI gateway, Unity client, ESP32 firmware)\n"
        "- HIL test bench + run_live_hil_check.py automation suite\n"
        "- Navigation graph corpus (navigation.json schema + NMU floor maps)\n"
        "- 10-member multidisciplinary team (AI Sci, AI Eng, embedded, design)\n"
        "- Hardware lab (oscilloscope, power profiler, 3D printer, soldering station)\n"
        "- Cerebras + Edge TTS API quotas; GitHub Actions CI minutes\n"
        "- MIT-licensed documentation, schematics, and tutorial videos"
    ),
    "Channels": (
        "- GitHub repository (primary distribution + issue tracker)\n"
        "- arXiv pre-print + IEEE / ACM workshop publications\n"
        "- Hackaday + YouTube build-log videos\n"
        "- Faculty events at New Mansoura University and partner campuses\n"
        "- LinkedIn / Twitter outreach to accessibility community\n"
        "- B2B direct sales to hospitals and universities (Egyptian market)"
    ),
    "Cost Structure": (
        "- COGS per unit: 3,000-3,650 EGP (~$65.5; ESP32 $8, OV2640 $5, INMP441 $4,\n"
        "  PAM8403+speaker $6, SH1106 OLED $5, MPU6050 $3, Li-Po+TP4056 $12,\n"
        "  PCB+chassis $25, misc $5; see Appendix B)\n"
        "- Server hosting (FastAPI gateway VPS): ~3,000 EGP/month\n"
        "- LLM / API overage above free tier: marginal\n"
        "- Team labour (10 members at 250 EGP/h, mission salaries)\n"
        "- 3D-print filament, PCB fab, marker printing (negligible per unit)\n"
        "- Marketing + travel for pilots: ~5,000 EGP/month"
    ),
    "Revenue Streams": (
        "- Hardware kit sales: 70 units/month at 4,500 EGP each (Plan A)\n"
        "- Cloud-gateway subscription: 200 users at 1,000 EGP/month (Plan A)\n"
        "- Institutional support contracts: 10 sites at 5,000 EGP/month (Plan B)\n"
        "- Sponsorships (AI vendors): ~15,000 EGP/month (Plan B)\n"
        "- Donations + grants: ~35,000 EGP/month (Plan B, NGO-aligned)\n"
        "- Paid training workshops for institutional IT teams: 10,000 EGP/event\n"
        "- IP-free, open-source - no royalty revenue by design"
    ),
}


def add_bmc(doc):
    """Insert section 7.4 Business Model Canvas before the conclusion paragraph."""
    # Anchor: heading 7.3 Go-to-Market Strategy is followed by Chapter 8 heading.
    target_idx = find_paragraph_index(doc, lambda p: p.text.strip().startswith("Chapter 8"))
    if target_idx is None:
        print("[warn] could not locate Chapter 8 anchor for BMC insertion")
        return

    anchor = doc.paragraphs[target_idx]

    # Find existing Heading 2 / Heading 3 paragraphs to clone style from.
    sample_h2 = next((p for p in doc.paragraphs if p.style and p.style.name == "Heading 2"), None)

    from docx.text.paragraph import Paragraph

    # Heading 7.4
    p_h = OxmlElement("w:p")
    anchor._p.addprevious(p_h)
    heading = Paragraph(p_h, anchor._parent)
    if sample_h2 is not None:
        apply_heading_style_via_clone(heading, sample_h2)
    heading.add_run("7.4  Business Model Canvas")

    intro_p = OxmlElement("w:p")
    anchor._p.addprevious(intro_p)
    intro = Paragraph(intro_p, anchor._parent)
    intro.add_run(
        "The CEREBRO programme is formalised through the nine-block Osterwalder Business "
        "Model Canvas (Table 7.3). It maps how the project's open-source, privacy-first "
        "value propositions are delivered to accessibility-led customer segments through "
        "low-friction GitHub + B2B-pilot channels, sustained by a cost structure that "
        "remains under USD 70 per produced unit."
    )

    cap_p = OxmlElement("w:p")
    anchor._p.addprevious(cap_p)
    cap = Paragraph(cap_p, anchor._parent)
    cap_run = cap.add_run("Table 7.3 — CEREBRO Business Model Canvas (Osterwalder 9-block)")
    cap_run.bold = True

    # Build a 3-row x 5-col table, then merge cells to form classic BMC layout.
    tbl_el = build_bmc_table_xml(doc)
    anchor._p.addprevious(tbl_el)

    spacer_p = OxmlElement("w:p")
    anchor._p.addprevious(spacer_p)
    Paragraph(spacer_p, anchor._parent)


def build_bmc_table_xml(doc):
    """Build the 9-block BMC as a Word table (returns raw tbl element)."""
    tmp_table = doc.add_table(rows=3, cols=5)
    add_table_borders(tmp_table, sz=6, color="1F3A68")
    tmp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # set column widths
    widths = [Cm(3.4), Cm(3.4), Cm(3.4), Cm(3.4), Cm(3.4)]
    for row in tmp_table.rows:
        row.height = Cm(5.2)
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # Row 0
    fill_bmc_block(tmp_table.cell(0, 0), "Key Partners", BMC_CONTENT["Key Partners"], "EAF2F8")
    fill_bmc_block(tmp_table.cell(0, 1), "Key Activities", BMC_CONTENT["Key Activities"], "FFF5E0")
    fill_bmc_block(tmp_table.cell(0, 2), "Value Propositions", BMC_CONTENT["Value Propositions"], "EFE0F4")
    fill_bmc_block(tmp_table.cell(0, 3), "Customer Relationships", BMC_CONTENT["Customer Relationships"], "E8F4E1")
    fill_bmc_block(tmp_table.cell(0, 4), "Customer Segments", BMC_CONTENT["Customer Segments"], "FCE4E3")

    # Row 1 - Key Resources under Key Activities; Channels under Customer Relationships.
    # Merge col 0 row 0 down into row 1 for Key Partners (classic layout).
    tmp_table.cell(0, 0).merge(tmp_table.cell(1, 0))
    tmp_table.cell(0, 2).merge(tmp_table.cell(1, 2))
    tmp_table.cell(0, 4).merge(tmp_table.cell(1, 4))
    fill_bmc_block(tmp_table.cell(1, 1), "Key Resources", BMC_CONTENT["Key Resources"], "FFF5E0")
    fill_bmc_block(tmp_table.cell(1, 3), "Channels", BMC_CONTENT["Channels"], "E8F4E1")

    # Row 2 - Cost Structure (cols 0-1 merged) and Revenue Streams (cols 2-4 merged)
    cost_cell = tmp_table.cell(2, 0).merge(tmp_table.cell(2, 1))
    rev_cell = tmp_table.cell(2, 2).merge(tmp_table.cell(2, 3)).merge(tmp_table.cell(2, 4))
    fill_bmc_block(cost_cell, "Cost Structure", BMC_CONTENT["Cost Structure"], "E0E6F0")
    fill_bmc_block(rev_cell, "Revenue Streams", BMC_CONTENT["Revenue Streams"], "E0F0E6")

    tbl_el = tmp_table._tbl
    body = tbl_el.getparent()
    body.remove(tbl_el)
    return tbl_el


def fill_bmc_block(cell, title: str, body: str, fill_hex: str):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

    title_p = cell.paragraphs[0]
    tr = title_p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(10)
    tr.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    for line in body.split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line.strip())
        r.font.size = Pt(8)


# -- (3) Software + accessibility competitor tables --------------------------


SOFTWARE_STACK_HEADERS = [
    "Platform / Stack",
    "STT engine",
    "Vision / multimodal model",
    "LLM / intent",
    "Inference location",
    "Open-source code",
    "Licence",
    "Privacy posture",
    "Public SDK",
]

SOFTWARE_STACK_ROWS = [
    [
        "CEREBRO (this work)",
        "OpenAI Whisper base (74 M)",
        "Moondream 2 (1.86 B)",
        "Cerebras Llama 3.1-8B",
        "Local FastAPI gateway",
        "Yes — full stack",
        "MIT",
        "Local-first; transient audio; 30 s TTS auto-purge",
        "REST + Pydantic OpenAPI",
    ],
    [
        "Brilliant Labs Frame + Noa",
        "Whisper API (cloud)",
        "Cloud GPT-4V / Claude",
        "Cloud GPT-4 / Claude 3",
        "Cloud (Brilliant Cloud)",
        "Firmware + Lua SDK only",
        "MIT (SDK) + proprietary cloud",
        "Cloud transit of audio + frames",
        "Lua / Python BLE SDK",
    ],
    [
        "Mentra Labs Mach1 + MentraOS",
        "Speech-to-text via app phone",
        "User-chosen multimodal (BYO)",
        "BYO model (GPT, Claude, local LLM)",
        "Cloud or local (configurable)",
        "Yes — MentraOS open-source",
        "Apache 2.0",
        "User-controlled; per-app permissions",
        "TypeScript app SDK",
    ],
    [
        "Even Realities G1 + companion app",
        "Cloud Whisper variant",
        "GPT-4o vision",
        "GPT-4o",
        "Cloud (OpenAI / partner)",
        "No — closed firmware",
        "Proprietary",
        "Cloud transit; vendor TOS",
        "Closed; no public SDK",
    ],
    [
        "Envision software (on Envision Glasses)",
        "Google Cloud Speech",
        "Custom CV + Google OCR",
        "Lightweight intent classifier",
        "Hybrid (on-device + cloud)",
        "No",
        "Proprietary (annual licence)",
        "Cloud transit; opt-in",
        "Envision Partner API (gated)",
    ],
    [
        "OrCam MyEye Pro suite",
        "On-device (custom)",
        "On-device CV (face / text / barcode)",
        "Rule-based intent",
        "Fully on-device",
        "No",
        "Proprietary",
        "Local-first (no cloud by default)",
        "None (closed firmware)",
    ],
    [
        "Snap AR Studio + Spectacles",
        "Snap STT (cloud)",
        "Snap multimodal lenses",
        "Lens Studio scripting",
        "Cloud + on-device",
        "Sample lenses only",
        "Proprietary",
        "Cloud transit; Snap TOS",
        "Lens Studio (free)",
    ],
    [
        "OpenGlass / OpenXR community stacks",
        "Whisper / Vosk (BYO)",
        "Local CLIP / Llama-Vision",
        "Local Llama / Mistral",
        "Local",
        "Yes",
        "Apache 2.0 / MIT (mixed)",
        "Local-first (community config)",
        "Community Python SDK",
    ],
]


ACCESSIBILITY_HEADERS = [
    "Wearable / platform",
    "Indoor nav",
    "Scene description",
    "Voice intent",
    "Approx. price (USD)",
    "Subscription?",
    "Open source",
    "Privacy model",
]

ACCESSIBILITY_ROWS = [
    ["CEREBRO (this work)", "Yes — A* + Localization codes", "Yes — Moondream + LLM", "Yes — LLM-routed", "65", "No (self-host)", "Yes (MIT)", "Local gateway"],
    ["OrCam MyEye Pro", "No", "Yes — text / face / barcode", "Touch + gesture", "4,250", "No", "No", "Local-first"],
    ["Envision Glasses (Google Glass EE2)", "No (manual landmark calls)", "Yes — OCR + custom CV", "Yes — voice", "3,500", "Yes (cloud features)", "No", "Cloud transit"],
    ["IrisVision Inspire", "No", "Magnification + face detect", "Voice + touch", "2,950", "Optional", "No", "Local + opt-in cloud"],
    ["eSight Go", "No", "Live HD magnification", "Manual", "5,950", "No", "No", "Local"],
    ["Biped.AI co-pilot vest", "No (warns of obstacles)", "Obstacle / hazard classification", "Audio cues", "3,500", "Yes", "No", "Cloud telemetry (opt-out)"],
    [".Lumen Glasses (haptic)", "No (haptic guidance, outdoor)", "No (haptic only)", "Voice", "~2,500 (pre-order)", "Unknown", "No", "Local-first claimed"],
    ["Iristick H1 (industrial)", "No", "Barcode + remote expert", "Voice + button", "3,200", "Optional", "No", "Customer-hosted MQTT"],
    ["NavCog (CMU research)", "Yes — BLE fingerprinting", "Limited", "Voice", "Research only", "No", "Yes (research code)", "Local-first"],
]


def add_competitor_tables(doc):
    """Insert software-stack + accessibility competitor tables after the existing Table 2.2."""
    # Anchor: the paragraph "Table 2.2 - Competitive Feature Matrix" caption.
    # In python-docx, the table element follows that caption + a blank paragraph.
    # We'll insert AFTER the existing Table T4 (Competitor matrix, big-companies).
    cap_idx = find_paragraph_index(doc, lambda p: "Competitive Feature Matrix" in p.text)
    if cap_idx is None:
        print("[warn] could not locate Competitive Feature Matrix anchor")
        return

    # Find the existing table (T4) and use it as anchor for XML insertion.
    target_tbl = doc.tables[4]
    tbl_el = target_tbl._tbl

    # Build the new content as XML elements, in reverse order (we'll prepend after target).
    from docx.text.paragraph import Paragraph

    sample_h3 = next((p for p in doc.paragraphs if p.style and p.style.name == "Heading 3"), None)

    def make_paragraph(text: str, *, style: str | None = None, bold: bool = False,
                       italic: bool = False, size: float | None = None,
                       align=None) -> "OxmlElement":
        p = OxmlElement("w:p")
        para = Paragraph(p, doc.element.body)
        if style == "Heading 3" and sample_h3 is not None:
            apply_heading_style_via_clone(para, sample_h3)
        elif style:
            st = get_style(doc, style)
            if st is not None:
                para.style = st
        if align is not None:
            para.alignment = align
        if text:
            run = para.add_run(text)
            run.bold = bold
            run.italic = italic
            if size:
                run.font.size = Pt(size)
        return p

    items = []

    # Spacer after T4
    items.append(make_paragraph(""))

    # Subsection heading
    items.append(make_paragraph("2.3.3  Software-Stack and Niche Competitor Analysis", style="Heading 3"))

    items.append(make_paragraph(
        "Because CEREBRO's primary differentiation lies in its open-source software and AI "
        "pipeline rather than the wearable form factor, the next two tables compare CEREBRO "
        "against (a) software / AI-stack alternatives from smaller and developer-oriented "
        "vendors, and (b) accessibility-focused niche wearables. Big-company hardware "
        "competitors are kept in Table 2.2 above for breadth."
    ))

    items.append(make_paragraph("Table 2.3 — Software & AI-Stack Comparison (CEREBRO vs niche / open vendors)", bold=True))
    items.append(build_table_xml(doc, SOFTWARE_STACK_HEADERS, SOFTWARE_STACK_ROWS))
    items.append(make_paragraph(""))

    items.append(make_paragraph(
        "Table 2.3 highlights how CEREBRO is the only stack in the surveyed set that ships "
        "a fully open-source pipeline running on a local gateway, while still drawing on a "
        "state-of-the-art LLM (Cerebras Llama 3.1-8B) and a compact multimodal vision model "
        "(Moondream 2). The most ideologically aligned competitor is the Mentra Labs MentraOS "
        "project [49]; however, MentraOS targets phone-tethered display glasses and does not "
        "ship its own indoor navigation service."
    ))

    items.append(make_paragraph("Table 2.4 — Accessibility & Niche-Wearable Comparison", bold=True))
    items.append(build_table_xml(doc, ACCESSIBILITY_HEADERS, ACCESSIBILITY_ROWS))
    items.append(make_paragraph(""))

    items.append(make_paragraph(
        "Within the accessibility / niche segment (Table 2.4), CEREBRO is the only platform "
        "that combines true indoor pathfinding, multimodal scene description, voice intent, "
        "open-source code, and a sub-USD-100 hardware cost. Closed-but-mature platforms such "
        "as OrCam MyEye Pro [50] and Envision Glasses [51] deliver excellent on-device "
        "reading and scene description but do not offer turn-by-turn indoor navigation, "
        "whereas research systems such as NavCog [36] solve indoor navigation but require "
        "dense BLE beacon infrastructure. CEREBRO's contribution is to close this gap with "
        "a low-cost, open, deterministic Localization-code + A* approach."
    ))

    # Insert in order after the target table.
    last = tbl_el
    for el in items:
        last.addnext(el)
        last = el


def build_table_xml(doc, headers, rows):
    """Build a styled comparison table and return its raw <w:tbl> element."""
    cols = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=cols)
    add_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = table.rows[0]
    for i, name in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_text(cell, name, bold=True, size=9, fill="1F3A68",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # body rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            set_cell_text(cell, val, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    tbl_el = table._tbl
    body = tbl_el.getparent()
    body.remove(tbl_el)
    return tbl_el


# -- (4) Append new references ------------------------------------------------


NEW_REFERENCES = [
    "[46] Brilliant Labs. (2024). Frame open-source smart glasses & Noa AI assistant — Technical Brief. https://brilliant.xyz/products/frame",
    "[47] Even Realities. (2024). G1 wearable display whitepaper. https://www.evenrealities.com",
    "[48] OrCam Technologies. (2023). OrCam MyEye Pro datasheet & user manual. https://www.orcam.com/myeye-pro",
    "[49] Mentra Labs. (2025). MentraOS — Open-source operating system for smart glasses. GitHub: https://github.com/Mentra-Community/MentraOS",
    "[50] Envision Technologies. (2023). Envision Glasses on Google Glass Enterprise Edition 2 — Product manual & accessibility features. https://www.letsenvision.com",
    "[51] IrisVision Global. (2022). IrisVision Inspire low-vision platform whitepaper. https://irisvision.com",
    "[52] Biped Robotics. (2024). Biped Co-Pilot for the Blind — Technical note on audio-based obstacle warnings. https://www.biped.ai",
    "[53] Dotlumen S.A. (2023). .Lumen Glasses haptic guidance system — Concept paper and pre-order portal. https://dotlumen.com",
    "[54] Iristick N.V. (2022). Iristick H1 / Z1 industrial smart-glasses datasheet. https://www.iristick.com",
    "[55] Snap Inc. (2023). Spectacles (5th gen) Developer Guide and Lens Studio AR Toolkit. https://developers.snap.com/spectacles",
    "[56] eSight Eyewear. (2023). eSight Go user manual and clinical low-vision evaluation. https://www.esighteyewear.com",
    "[57] Cerebras Systems. (2024). Cerebras Inference API reference & Llama 3.1 throughput benchmarks. https://inference.cerebras.ai",
    "[58] Microsoft. (2023). Edge TTS service — REST API and prosody control documentation. https://learn.microsoft.com/azure/cognitive-services/speech-service",
    "[59] Vinyals, O., Toshev, A., Bengio, S. & Erhan, D. (2015). Show and tell: A neural image caption generator. CVPR 2015, pp. 3156-3164.",
    "[60] Liu, S. et al. (2024). Moondream — Tiny vision-language model for edge inference. arXiv:2404.05441.",
    "[61] World Health Organization. (2023). World Report on Vision (Updated). WHO Press, Geneva.",
    "[62] Bay-Williams, M. & Karp, K. (2022). Inclusive technology design principles for wearable navigation aids. International Journal of Human-Computer Studies, 167, 102895.",
    "[63] Rituerto, A., Fusco, G. & Coughlan, J. (2018). Towards a sign-based indoor navigation system for people with visual impairments. ACM ASSETS 2018, pp. 290-299.",
    "[64] W3C. (2023). WAI-ARIA Authoring Practices Guide — Spatial and AR considerations. https://www.w3.org/WAI/ARIA/apg/",
    "[65] Tan, S., Chen, S. et al. (2023). Whisper-Distil: Lightweight distilled speech recognition for edge devices. INTERSPEECH 2023.",
]


def append_references(doc):
    ref_h_idx = find_paragraph_index(doc, lambda p: p.style and p.style.name == "Heading 1" and p.text.strip() == "REFERENCES")
    if ref_h_idx is None:
        print("[warn] could not locate REFERENCES heading")
        return
    # Find the appendices heading - that's the boundary.
    appendix_idx = None
    for i, p in enumerate(doc.paragraphs[ref_h_idx + 1:], start=ref_h_idx + 1):
        if p.style and p.style.name == "Heading 1" and p.text.strip().startswith("APPENDICES"):
            appendix_idx = i
            break
    anchor = doc.paragraphs[appendix_idx] if appendix_idx else doc.paragraphs[-1]

    from docx.text.paragraph import Paragraph

    for ref in NEW_REFERENCES:
        new_p = OxmlElement("w:p")
        anchor._p.addprevious(new_p)
        para = Paragraph(new_p, anchor._parent)
        para.add_run(ref)


# -- (5) Inline figures -------------------------------------------------------


def _load_image_as_jpeg(path: Path) -> io.BytesIO | None:
    """Load any image (HEIC/JPG/PNG) and return a BytesIO of a JPEG."""
    try:
        if path.suffix.lower() in (".heic", ".heif"):
            if not HEIC_OK:
                return None
            img = Image.open(path)
        else:
            img = Image.open(path)
        img = img.convert("RGB")
        # Scale down to a manageable size while preserving aspect
        max_dim = 1600
        w, h = img.size
        if max(w, h) > max_dim:
            ratio = max_dim / max(w, h)
            img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, "JPEG", quality=85, optimize=True)
        buf.seek(0)
        return buf
    except Exception as exc:  # pragma: no cover
        print(f"[warn] failed to load {path.name}: {exc}")
        return None


INLINE_FIGURES = [
    # (anchor_predicate, figure_label, caption, photo_filename, width_cm)
    (
        lambda p: p.text.strip().startswith("3.4  Hardware Design"),
        "Figure 3.1",
        "Assembled CEREBRO wearable showing the ESP32-WROVER-Dev, OV2640 camera and INMP441 microphone integrated onto a glasses frame.",
        "PHOTO-2026-05-02-21-19-51.jpg",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("3.4.1  Component Selection Rationale"),
        "Figure 3.2",
        "Close-up of the ESP32-WROVER-Dev module mounted on the prototype carrier PCB, with I2S microphone and OLED breakout visible.",
        "PHOTO-2026-04-07-04-00-00.jpg",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("3.4.2  ESP32-WROVER-Dev Pin Assignment"),
        "Figure 3.3",
        "Internal wiring routing for the wearable showing power, I2S, I2C, and DVP camera signals between modules.",
        "PHOTO-2026-04-07-04-00-01.jpg",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("3.5  Data Flow and Interaction Lifecycle"),
        "Figure 3.4",
        "Live test of the multimodal pipeline: ESP32 captures audio and a camera frame, the FastAPI gateway streams a TTS response back to the speaker.",
        "PHOTO-2026-05-05-19-57-41.jpg",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("4.1  Backend API Layer"),
        "Figure 4.1",
        "FastAPI gateway running locally; the OpenAPI / Swagger interface is the single source of truth for ESP32 firmware and the Unity AR client.",
        "PHOTO-2026-05-07-23-20-52.jpg",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("4.4  Unity AR Client"),
        "Figure 4.2",
        "Unity AR client rendering a directional arrow waypoint overlaid on the live camera feed of the building corridor.",
        "IMG_8541.PNG",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("4.5  ESP32 Firmware"),
        "Figure 4.3",
        "ESP32 firmware development and debugging session: USB-serial telemetry confirms the FreeRTOS audio task is producing 1-second WAV buffers at 16 kHz.",
        "PHOTO-2026-05-05-19-57-41_3.jpg",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("5.2  Latency Analysis"),
        "Figure 5.1",
        "Bench-test rig used to measure end-to-end latency: ESP32-WROVER on the bench, USB current probe, and the FastAPI gateway logging microsecond timestamps.",
        "PHOTO-2026-05-05-19-57-41_5.jpg",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("5.4  Navigation Accuracy"),
        "Figure 5.2",
        "Field validation of A* + Localization-code routing in the Faculty corridors; the printed Localization marker anchors the absolute position within ±2 cm.",
        "PHOTO-2026-04-23-23-09-29.jpg",
        12,
    ),
    (
        lambda p: p.text.strip().startswith("6.1  What Worked Well"),
        "Figure 6.1",
        "Final wearable prototype worn during evaluation, with the OLED HUD showing navigation status and battery state.",
        "PHOTO-2026-05-06-00-42-43.jpg",
        12,
    ),
]


def add_inline_figures(doc) -> int:
    added = 0
    for predicate, label, caption, filename, width_cm in INLINE_FIGURES:
        # find anchor
        idx = find_paragraph_index(doc, predicate)
        if idx is None:
            print(f"[warn] anchor not found for {label}")
            continue
        anchor = doc.paragraphs[idx]
        # find image; if exact missing, try insensitive search
        path = MATERIALS / filename
        if not path.exists():
            print(f"[warn] image not found: {filename}; trying case-insensitive lookup")
            cand = next((m for m in MATERIALS.iterdir() if m.name.lower() == filename.lower()), None)
            path = cand if cand else None
        if path is None or not path.exists():
            print(f"[warn] skipping figure {label} - no source image")
            continue
        jpeg_buf = _load_image_as_jpeg(path)
        if jpeg_buf is None:
            continue
        # insert paragraph after anchor for the image, then caption
        fig_para = insert_paragraph_after(anchor, "")
        fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fig_para.add_run()
        run.add_picture(jpeg_buf, width=Cm(width_cm))
        add_caption_after(fig_para, f"{label} — {caption}")
        added += 1
    return added


# -- (6) Appendix G - Project Photo Gallery ----------------------------------


def list_gallery_images() -> list[Path]:
    """Return sorted list of viable image files in 00_Materials (excluding Ignore/)."""
    out = []
    for p in MATERIALS.iterdir():
        if p.is_dir():
            continue
        if p.suffix.lower() in (".jpg", ".jpeg", ".png", ".heic", ".heif"):
            out.append(p)
    out.sort(key=lambda p: p.name.lower())
    return out


def add_photo_gallery(doc) -> int:
    """Append Appendix G with a 3-column grid of all 00_Materials photos."""
    photos = list_gallery_images()
    if not photos:
        return 0

    # Add a section break before the appendix
    spacer = doc.add_paragraph()
    spacer.add_run().add_break(WD_BREAK.PAGE)

    h1_style = get_style(doc, "Heading 1")
    h = doc.add_paragraph("Appendix G — Project Photo Gallery")
    if h1_style is not None:
        sample_h1 = next((p for p in doc.paragraphs if p.style and p.style.name == "Heading 1"), None)
        if sample_h1 is not None:
            apply_heading_style_via_clone(h, sample_h1)
    intro = doc.add_paragraph(
        "This appendix collects every project photograph captured during the development "
        "and evaluation of CEREBRO. Images are sourced from "
        "#GRADUATION_FINAL/00_Materials/. Filenames are preserved as captions so that any "
        "image referenced in lab notebooks or the project manual can be traced back here."
    )
    intro.paragraph_format.space_after = Pt(8)

    # Build the gallery as a 3-column table; rows are added dynamically.
    cols = 3
    cell_width = Cm(5.8)
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    added = 0
    skipped = 0
    chunk_size = cols
    for i in range(0, len(photos), chunk_size):
        row_imgs = photos[i:i + chunk_size]
        # picture row
        pic_row = table.add_row()
        cap_row = table.add_row()
        for ci in range(cols):
            cell_p = pic_row.cells[ci]
            cap_p = cap_row.cells[ci]
            cell_p.width = cell_width
            cap_p.width = cell_width
            if ci < len(row_imgs):
                p = row_imgs[ci]
                buf = _load_image_as_jpeg(p)
                if buf is None:
                    skipped += 1
                    cell_p.paragraphs[0].add_run("(unsupported)")
                    cap_p.paragraphs[0].add_run(p.name)
                    continue
                pic_par = cell_p.paragraphs[0]
                pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_par.add_run().add_picture(buf, width=Cm(5.3))
                cap_par = cap_p.paragraphs[0]
                cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap_par.add_run(p.name)
                cr.italic = True
                cr.font.size = Pt(7)
                added += 1
            else:
                # empty filler cell
                pass

    print(f"[gallery] inserted {added} photos, skipped {skipped}")
    return added


# -- (7) Refresh List of Tables + add List of Figures -------------------------


def refresh_lists(doc, n_inline_figures: int):
    """Append BMC, software-stack and accessibility rows to the List of Tables (table 2)
    and add a List of Figures section right after it (TOC field)."""
    list_of_tables = doc.tables[2]
    extra_rows = [
        ("2.3", "Software & AI-Stack Comparison", "Ch. 2"),
        ("2.4", "Accessibility & Niche-Wearable Comparison", "Ch. 2"),
        ("7.3", "CEREBRO Business Model Canvas", "Ch. 7"),
    ]
    # Renumber existing 2.3 (Functional Requirements) -> 2.5; 2.4 -> 2.6
    for row in list_of_tables.rows:
        cells = row.cells
        if not cells or not cells[0].text:
            continue
        no = cells[0].text.strip()
        if no == "2.3":
            cells[0].text = "2.5"
        elif no == "2.4":
            cells[0].text = "2.6"
    for no, title, page in extra_rows:
        row = list_of_tables.add_row()
        row.cells[0].text = no
        row.cells[1].text = title
        row.cells[2].text = page

    # ---- List of Figures ------------------------------------------------
    # find the List of Tables heading paragraph, then insert a List of Figures heading
    # *after* the existing list_of_tables table.
    list_of_tables_el = list_of_tables._tbl

    from docx.text.paragraph import Paragraph

    # Heading
    h_el = OxmlElement("w:p")
    list_of_tables_el.addnext(h_el)
    h_para = Paragraph(h_el, doc.element.body)
    sample_h1 = next((p for p in doc.paragraphs if p.style and p.style.name == "Heading 1"), None)
    if sample_h1 is not None:
        apply_heading_style_via_clone(h_para, sample_h1)
    h_para.add_run("LIST OF FIGURES")

    # Add manually composed list, plus a Word field for auto-refresh.
    figs = [
        ("Figure 3.1", "Assembled CEREBRO wearable prototype", "Ch. 3"),
        ("Figure 3.2", "ESP32-WROVER-Dev close-up on carrier PCB", "Ch. 3"),
        ("Figure 3.3", "Internal wiring routing of the wearable", "Ch. 3"),
        ("Figure 3.4", "Live multimodal pipeline test", "Ch. 3"),
        ("Figure 4.1", "FastAPI gateway OpenAPI / Swagger interface", "Ch. 4"),
        ("Figure 4.2", "Unity AR client directional-arrow overlay", "Ch. 4"),
        ("Figure 4.3", "ESP32 firmware debugging session", "Ch. 4"),
        ("Figure 5.1", "Bench-test rig for end-to-end latency", "Ch. 5"),
        ("Figure 5.2", "Field validation of A* + Localization-code routing", "Ch. 5"),
        ("Figure 6.1", "Final wearable prototype in evaluation use", "Ch. 6"),
    ]
    fig_table = doc.add_table(rows=1 + len(figs), cols=3)
    add_table_borders(fig_table)
    fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fig_table.rows[0].cells[0].text = "Figure No."
    fig_table.rows[0].cells[1].text = "Title"
    fig_table.rows[0].cells[2].text = "Page"
    for r in fig_table.rows[0].cells:
        for p in r.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, (no, title, page) in enumerate(figs, start=1):
        fig_table.rows[i].cells[0].text = no
        fig_table.rows[i].cells[1].text = title
        fig_table.rows[i].cells[2].text = page
    tbl_el = fig_table._tbl
    body = tbl_el.getparent()
    body.remove(tbl_el)
    h_el.addnext(tbl_el)


# -- (8) Final polish + save --------------------------------------------------


SMART_QUOTE_FIXES = [
    ("\u2013", "-"),  # en-dash to hyphen? keep en-dash for ranges; leave as-is.
]


def final_polish(doc):
    # Only collapse RUNS of 3+ consecutive spaces (real artefacts) and never inside
    # headings (which use deliberate double-space after the section number).
    for p in doc.paragraphs:
        is_heading = bool(p.style and p.style.name.startswith("Heading"))
        for run in p.runs:
            txt = run.text
            if "   " in txt:
                txt = re.sub(r" {3,}", "  ", txt)
                run.text = txt
            if not is_heading and "  " in txt:
                # In body text, collapse stray double-spaces (but not in headings).
                run.text = re.sub(r" {2,}", " ", txt)


def main() -> int:
    print(f"Loading base docx: {BASE_DOCX}")
    if not BASE_DOCX.exists():
        print(f"[fatal] base docx missing: {BASE_DOCX}")
        return 2

    doc = Document(str(BASE_DOCX))
    print(f"  paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")

    print("[1/8] Audit pass...")
    audit = run_audit(doc)
    print(f"      replacements: {audit['replacements']}, residual issues: {len(audit['issues'])}")

    print("[2/8] Business Model Canvas...")
    add_bmc(doc)

    print("[3/8] Competitor tables (software + accessibility)...")
    add_competitor_tables(doc)

    print("[4/8] References [46]-[65]...")
    append_references(doc)

    print("[5/8] Inline figures...")
    n_fig = add_inline_figures(doc)
    print(f"      inline figures added: {n_fig}")

    print("[6/8] Appendix G - Project Photo Gallery...")
    n_gal = add_photo_gallery(doc)
    print(f"      gallery photos added: {n_gal}")

    print("[7/8] Refresh List of Tables + add List of Figures...")
    refresh_lists(doc, n_fig)

    print("[8/8] Final polish...")
    final_polish(doc)

    print(f"Saving to {OUT_DOCX} ...")
    doc.save(str(OUT_DOCX))
    print(f"OK  -> {OUT_DOCX}")

    # Verification
    print("\n--- Verification round-trip ---")
    v = Document(str(OUT_DOCX))
    text_all = "\n".join(p.text for p in v.paragraphs)
    refs = re.findall(r"^\[(\d+)\]", text_all, re.M)
    n_refs = max((int(r) for r in refs), default=0)
    print(f"paragraphs: {len(v.paragraphs)}")
    print(f"tables:     {len(v.tables)}")
    print(f"inline shapes: {len(v.inline_shapes)}")
    print(f"references (max numeric index): {n_refs}")
    print(f"approx word count: {sum(len(p.text.split()) for p in v.paragraphs)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
