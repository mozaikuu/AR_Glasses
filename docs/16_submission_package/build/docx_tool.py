#!/usr/bin/env python3
"""
Direct .docx tools for the graduation package (inspect, text replace, template fill, core props).

Install:
  pip install -r docs/16_submission_package/build/requirements-docx.txt

Examples:
  python docx_tool.py inspect docs/00_Materials/Graudation-project-template.docx
  python docx_tool.py fill-template docs/00_Materials/Graudation-project-template.docx \\
      docs/16_submission_package/metadata.yaml docs/16_submission_package/dist/Faculty_Template_Filled.docx
  python docx_tool.py set-core-props docs/16_submission_package/dist/SmartGlasses_ProjectManual.docx \\
      docs/16_submission_package/metadata.yaml
  python docx_tool.py replace docs/in.docx docs/out.docx "OLD" "NEW"
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from pathlib import Path
from typing import Iterable

try:
    import yaml
    from docx import Document
    from docx.document import Document as DocumentType
except ImportError as e:
    print("Missing dependency. Run:", file=sys.stderr)
    print("  pip install -r docs/16_submission_package/build/requirements-docx.txt", file=sys.stderr)
    raise SystemExit(1) from e


def _load_yaml(path: Path) -> dict:
    data = yaml.safe_load(path.read_text(encoding="utf-8"))
    return data if isinstance(data, dict) else {}


def iter_all_paragraphs(doc: DocumentType) -> Iterable:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr is None:
                continue
            for p in hdr.paragraphs:
                yield p
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    paragraph.text = paragraph.text.replace(old, new)
    return True


def replace_all(doc: DocumentType, old: str, new: str) -> int:
    n = 0
    for p in iter_all_paragraphs(doc):
        if replace_in_paragraph(p, old, new):
            n += 1
    return n


def replace_first_global(doc: DocumentType, old: str, new: str) -> bool:
    """Replace only the first occurrence of `old` in document order."""
    for p in iter_all_paragraphs(doc):
        if old in p.text:
            p.text = p.text.replace(old, new, 1)
            return True
    return False


def _insert_paragraph_after(paragraph, text: str) -> None:
    """Insert a new body paragraph immediately after ``paragraph`` (OOXML)."""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)


def _append_third_team_line(doc: DocumentType, line: str, second_id: str) -> None:
    """Template has two name/ID slots; add a third line after the second student if needed."""
    anchor = None
    for p in doc.paragraphs:
        if second_id and second_id in p.text:
            anchor = p
            break
    if anchor is not None and line:
        _insert_paragraph_after(anchor, line)


def _safe_print(s: str) -> None:
    enc = getattr(sys.stdout, "encoding", None) or "utf-8"
    try:
        print(s)
    except UnicodeEncodeError:
        print(s.encode(enc, errors="replace").decode(enc, errors="replace"))


def cmd_inspect(args: argparse.Namespace) -> int:
    path = Path(args.docx)
    doc = Document(str(path))
    _safe_print(f"File: {path.resolve()}")
    _safe_print(f"Paragraphs (body): {len(doc.paragraphs)}")
    _safe_print(f"Tables (body): {len(doc.tables)}")
    _safe_print("--- First 40 non-empty paragraph texts (trimmed) ---")
    shown = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        snippet = t[:200].encode("ascii", errors="replace").decode("ascii")
        _safe_print(f"[{shown}] {snippet!r}")
        shown += 1
        if shown >= 40:
            break
    _safe_print("--- Core properties ---")
    cp = doc.core_properties
    _safe_print(" title: " + repr(cp.title))
    _safe_print(" subject: " + repr(cp.subject))
    _safe_print(" author: " + repr(cp.author))
    return 0


def cmd_replace(args: argparse.Namespace) -> int:
    src = Path(args.src)
    dst = Path(args.dst)
    shutil.copy2(src, dst)
    doc = Document(str(dst))
    count = replace_all(doc, args.old, args.new)
    doc.save(str(dst))
    _safe_print(f"Replaced {count!r} paragraph/cell blocks touching {args.old!r} -> wrote {dst}")
    return 0


def _authors_from_meta(meta: dict) -> tuple[list[str], list[str]]:
    """Return (names, ids) from metadata author list strings."""
    names: list[str] = []
    ids: list[str] = []
    raw = meta.get("author")
    if not isinstance(raw, list):
        raw = [raw] if raw else []
    for line in raw:
        s = str(line).strip()
        m = re.search(r"\(Student ID:\s*([^)]+)\)", s, re.I)
        if m:
            ids.append(m.group(1).strip())
            names.append(re.sub(r"\s*\(Student ID:.*\)\s*$", "", s, flags=re.I).strip())
        else:
            names.append(s)
            ids.append("")
    return names, ids


def cmd_fill_template(args: argparse.Namespace) -> int:
    template = Path(args.template)
    meta_path = Path(args.metadata)
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, out)
    meta = _load_yaml(meta_path)
    title = str(meta.get("title") or "Smart Glasses Distilled").strip()
    short_title = title.split("\u2014")[0].strip() if "\u2014" in title else title[:120]
    names, ids = _authors_from_meta(meta)

    doc = Document(str(out))
    # Order matters: `<Academic TITLE OF THE GRADUATION PROJECT>` contains substring
    # `TITLE OF THE GRADUATION PROJECT`; replace the long form first.
    replace_all(doc, "<Academic TITLE OF THE GRADUATION PROJECT>", title)
    replace_all(doc, "TITLE OF THE GRADUATION PROJECT", short_title)

    # Cover uses "Name - ID" on one line (see docx_tool.py inspect output).
    pair_ph = "<Name LastName> -<Student ID Number>"
    pair_ph_alt = "<Name LastName> - <Student ID Number>"
    for i in range(max(len(names), len(ids))):
        name = names[i] if i < len(names) else ""
        sid = ids[i] if i < len(ids) else ""
        line = f"{name} - {sid}".strip()
        if not replace_first_global(doc, pair_ph, line):
            replace_first_global(doc, pair_ph_alt, line)

    if len(names) > 2:
        third_line = (f"{names[2]} - {ids[2]}" if len(ids) > 2 and ids[2] else names[2]).strip()
        second_id = ids[1] if len(ids) > 1 else ""
        _append_third_team_line(doc, third_line, second_id)

    # Advisor line in template is a single paragraph: "<Title> <Name> <LastName>"
    replace_all(doc, "<Title> <Name> <LastName>", "(Advisor - official faculty assignment)")

    cp = doc.core_properties
    cp.title = title
    cp.author = "; ".join(n for n in names if n)
    if meta.get("institution"):
        cp.comments = str(meta["institution"])

    doc.save(str(out))
    _safe_print(f"Filled template -> {out.resolve()}")
    return 0


def cmd_set_core_props(args: argparse.Namespace) -> int:
    path = Path(args.docx)
    meta = _load_yaml(Path(args.metadata))
    title = str(meta.get("title") or "").strip()
    names, _ = _authors_from_meta(meta)
    doc = Document(str(path))
    cp = doc.core_properties
    if title:
        cp.title = title
    if names:
        cp.author = "; ".join(names)
    if meta.get("institution"):
        cp.comments = str(meta["institution"])
    doc.save(str(path))
    _safe_print(f"Updated core properties on {path.resolve()}")
    return 0


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Direct .docx utilities (python-docx).")
    sub = p.add_subparsers(dest="cmd", required=True)

    p_in = sub.add_parser("inspect", help="Print structure and sample paragraph text.")
    p_in.add_argument("docx", type=str)
    p_in.set_defaults(func=cmd_inspect)

    p_rep = sub.add_parser("replace", help="Copy src to dst and replace all occurrences of OLD with NEW.")
    p_rep.add_argument("src")
    p_rep.add_argument("dst")
    p_rep.add_argument("old")
    p_rep.add_argument("new")
    p_rep.set_defaults(func=cmd_replace)

    p_fill = sub.add_parser(
        "fill-template",
        help="Copy faculty Word template and replace title / student placeholders from metadata.yaml.",
    )
    p_fill.add_argument("template", help="Path to Graudation-project-template.docx")
    p_fill.add_argument("metadata", help="Path to metadata.yaml")
    p_fill.add_argument("out", help="Output .docx path")
    p_fill.set_defaults(func=cmd_fill_template)

    p_core = sub.add_parser(
        "set-core-props",
        help="Set Word core properties (title, author, comments) on an existing .docx in place.",
    )
    p_core.add_argument("docx")
    p_core.add_argument("metadata", help="Path to metadata.yaml")
    p_core.set_defaults(func=cmd_set_core_props)

    return p


def main() -> int:
    args = build_parser().parse_args()
    return int(args.func(args))


if __name__ == "__main__":
    raise SystemExit(main())
