from docx import Document

doc = Document()

# Title
doc.add_heading('Full Specification Table: Meta Ray-Ban Smart Glasses, Lenovo ThinkReality A3, Microsoft HoloLens 2', level=1)

# Define table data
headers = [
    "Category",
    "Meta Ray-Ban Smart Glasses",
    "Lenovo ThinkReality A3",
    "Microsoft HoloLens 2"
]

rows = [
    ["Device Type", "Smart glasses (camera + audio)", "AR smart glasses (tethered)", "Standalone mixed-reality headset"],
    ["Price", "~$299–$379", "~$1,499", "~$3,500"],
    ["Release Year", "2023 (2nd Gen)", "2021", "2019"],
    ["Weight", "~48–50 g", "~130 g", "~566 g"],
    ["Dimensions", "Eyewear form factor", "Glasses form factor", "Head-mounted visor"],
    ["Operating System", "Meta firmware", "Android-based custom OS", "Windows Holographic"],
    ["CPU / Processor", "Qualcomm Snapdragon (light SoC)", "Uses PC/phone CPU", "Qualcomm Snapdragon 850 + HPU 2.0"],
    ["GPU", "Integrated mobile GPU", "Uses PC/phone GPU", "Adreno + custom HPU"],
    ["RAM", "Not specified", "Uses PC/phone RAM", "4 GB LPDDR4x"],
    ["Storage", "Not applicable", "Uses PC/phone storage", "64 GB UFS"],
    ["Display Type", "None (not AR)", "1080p micro-OLED AR displays", "Waveguide holographic displays"],
    ["Display Resolution", "—", "1920 × 1080 per eye", "2048 × 1080 per eye"],
    ["Field of View (FOV)", "—", "~40° diagonal", "~52° diagonal"],
    ["Optics", "—", "45% transparency", "Diffractive waveguides"],
    ["Brightness", "—", "~200–400 nits", "~500 nits equivalent"],
    ["Cameras", "12 MP photo, 1080p60 video", "8-MP RGB + dual fisheye tracking", "Depth camera + 8-MP RGB"],
    ["Sensors", "Accelerometer, gyroscope, touch panel, mics", "IMU, tracking cameras", "IMU, depth sensor, IR cams, eye tracking"],
    ["Tracking Type", "No positional tracking", "6DoF inside-out", "6DoF inside-out"],
    ["Hand Tracking", "No", "Yes (basic)", "Yes (advanced)"],
    ["Eye Tracking", "No", "No", "Yes"],
    ["Environment Mapping", "No", "Spatial mapping", "Real-time spatial mapping"],
    ["Audio", "Open-ear speakers; 5-mic array", "Stereo speakers", "Spatial sound speakers"],
    ["Voice Assistant", "Meta AI", "No", "Cortana (enterprise optional)"],
    ["Connectivity", "Bluetooth 5.x", "USB-C tether", "Wi-Fi 5, Bluetooth"],
    ["Battery Life", "3–4 hours", "Tethered (no battery)", "2–3 hours"],
    ["Charging Time", "~1 hour", "—", "65–90 minutes"],
    ["Battery Capacity", "Not specified", "—", "~15.2 Wh"],
    ["Controller Support", "None", "Optional 6DoF controller", "Hand tracking only"],
    ["Compute Mode", "On-device", "Tethered", "Standalone"],
    ["Software Ecosystem", "Meta apps", "ThinkReality platform", "Dynamics 365, MR apps"],
    ["Development Platform", "Meta SDK", "ThinkReality SDK", "Unity, Unreal, MRTK"],
    ["Security / Privacy", "Capture LED, encrypted pairing", "Enterprise controls", "Enterprise security, secure boot"],
    ["Use Cases", "Social content, music, calls", "Enterprise AR workflows", "MR training, surgery, engineering"],
    ["Build Material", "Acetate frame", "Plastic/metal", "Carbon fiber, plastic"],
    ["Comfort / Fit", "Everyday sunglasses", "Light but tethered", "Balanced head-mounted"],
    ["Environmental Rating", "—", "Depends on PC/phone", "IP50"],
    ["Accessories", "Charging case", "Safety frame, PC/phone", "Fit kit, carry case"],
    ["Supported Platforms", "iOS, Android", "PC, Motorola phones", "Windows"],
    ["Developer Access", "Moderate", "Moderate", "Full enterprise SDK"],
    ["Safety Features", "Recording LED", "Enterprise admin controls", "Eye-safety compliance"],
]

# Create table
table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h

# Fill table rows
for row in rows:
    row_cells = table.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = item

# Save file
filepath = "./Smart_Glasses_Full_Specifications.docx"
doc.save(filepath)
